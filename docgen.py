from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os
import uuid

def generate_word_doc(questions_str, subject, grade):
    # 创建临时文件夹
    if not os.path.exists("./tmp"):
        os.makedirs("./tmp")
    
    doc = Document()
    
    # ==================== 设置全局样式 ====================
    # 设置正文字体（中文字体用宋体，西文Times New Roman）
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.size = Pt(12)  # 小四号字
    normal_style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    # ==================== 设置试卷标题 ====================
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{grade}{subject}考试试卷")
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)  # 一号字
    title_run.bold = True
    
    # 添加考试信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run("考试时间：120分钟　　满分：100分")
    info_run.font.name = '楷体'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    info_run.font.size = Pt(14)
    info.paragraph_format.space_after = Pt(24)  # 增加标题与正文间距
    
    # ==================== 设置页面布局 ====================
    section = doc.sections[0]
    section.top_margin = Cm(2.5)     # 上边距
    section.bottom_margin = Cm(2.5)  # 下边距
    section.left_margin = Cm(3)      # 左边距
    section.right_margin = Cm(3)     # 右边距
    
    # ==================== 添加试题内容 ====================
    lines = questions_str.strip().split("\n")
    current_question = ""
    current_question_type = None
    
    for line in lines:
        line = line.strip()
        
        # 检测题型标题（一、二、三、等）
        if line.startswith(('一、', '二、', '三、', '四、', '五、')):
            # 处理前一个问题（如果有）
            if current_question:
                add_question_paragraph(doc, current_question, current_question_type)
            
            # 添加题型标题
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)  # 题型前间距
            p.paragraph_format.space_after = Pt(12)   # 题型后间距
            run = p.add_run(line)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
            run.bold = True
            current_question_type = line.split('、')[0]
            current_question = ""
            continue
            
        # 合并连续的问题行（选项和答案）
        if line.startswith(("A.", "B.", "C.", "D.", "1.", "2.", "3.", "4.", "5.")) or line.endswith(('( )')):
            current_question += "\n" + line
        else:
            # 处理前一个问题（如果有）
            if current_question:
                add_question_paragraph(doc, current_question, current_question_type)
            
            # 开始新问题
            current_question = line
    
    # 处理最后一个问题
    if current_question:
        add_question_paragraph(doc, current_question, current_question_type)
    
    # ==================== 保存文件 ====================
    filename = f"./tmp/{uuid.uuid4()}.docx"
    doc.save(filename)
    return filename

def add_question_paragraph(doc, text, question_type=None):
    """添加问题段落，统一处理格式"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    
    # 设置段落格式
    p_format.space_before = Pt(6)  # 段前间距
    p_format.space_after = Pt(6)   # 段后间距
    p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 特殊处理选择题选项缩进
    if any(text.startswith(x) for x in ["A.", "B.", "C.", "D."]):
        p_format.left_indent = Cm(0.74)  # 选项缩进
    
    # 处理填空题的下划线
    if "_____" in text:
        run = p.add_run(text.replace("_____", "________"))
    else:
        run = p.add_run(text)
    
    # 设置字体
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 判断题的特殊格式
    if text.endswith(('( )')):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    

